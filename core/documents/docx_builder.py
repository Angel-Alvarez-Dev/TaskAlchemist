import os
from docx import Document
from datetime import datetime

class DocxBuilder:
    def __init__(self, tarea: dict):
        """
        Inicializa el generador de documentos con los datos de una tarea.

        Args:
            tarea (dict): Diccionario con la información de la tarea
        """
        self.tarea = tarea
        self.output_dir = os.getenv("DOCS_OUTPUT", "./output")

    def generate(self) -> str:
        """
        Genera un documento .docx con formato académico básico.

        Returns:
            str: Ruta del archivo generado
        """
        doc = Document()

        # Título de la tarea
        doc.add_heading(self.tarea['nombre'], level=1)

        # Información adicional
        doc.add_paragraph(f"Asignatura: {self.tarea['asignatura']}")
        fecha = datetime.fromisoformat(self.tarea['fecha_entrega']).strftime('%d/%m/%Y')
        doc.add_paragraph(f"Fecha de entrega: {fecha}")

        # Nota adicional si existe
        if self.tarea.get("nota"):
            doc.add_paragraph(f"Nota: {self.tarea['nota']}")

        # Espacio para desarrollo del contenido
        doc.add_paragraph("\nDesarrollo:")
        doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit.")

        # Ruta de salida organizada por asignatura
        asignatura_dir = os.path.join(self.output_dir, self.tarea['asignatura'])
        os.makedirs(asignatura_dir, exist_ok=True)

        file_name = f"{self.tarea['nombre'].replace(' ', '_')}.docx"
        file_path = os.path.join(asignatura_dir, file_name)

        doc.save(file_path)
        return file_path