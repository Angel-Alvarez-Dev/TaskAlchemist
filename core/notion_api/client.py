# core/notion_api/client.py
# core/notion_api/client.py
import subprocess
import win32com.client if os.name == 'nt' else None
from pathlib import Path
from docx.shared import Pt, Inches
from docx import Document
from datetime import datetime
from typing import List, Dict, Any
import os
from notion_client import Client


class NotionConnector:
    def __init__(self):
        self.client = Client(auth=os.getenv('NOTION_TOKEN'))
        self.database_id = os.getenv('NOTION_DATABASE_ID')

    def get_pending_tasks(self) -> List[Dict[str, Any]]:
        """Obtiene las tareas pendientes de la base de datos de Notion"""
        response = self.client.databases.query(
            database_id=self.database_id,
            filter={
                "and": [
                    {
                        "property": "Status",
                        "select": {
                            "equals": "Pendiente"
                        }
                    }
                ]
            }
        )
        return response.get('results', [])

    def mark_task_processed(self, task_id: str, file_path: str) -> None:
        """Actualiza el estado de la tarea y guarda la ruta del archivo"""
        self.client.pages.update(
            page_id=task_id,
            properties={
                "Status": {"select": {"name": "En Proceso"}},
                "FilePath": {"rich_text": [{"text": {"content": file_path}}]}
            }
        )

    def update_task_property(self, task_id: str, property_name: str, value: str) -> None:
        """Actualiza una propiedad específica de la tarea"""
        self.client.pages.update(
            page_id=task_id,
            properties={
                property_name: {"rich_text": [{"text": {"content": value}}]}
            }
        )


# core/document_manager/file_generator.py


class DocumentAlchemist:
    def __init__(self):
        self.base_path = os.getenv('DOCUMENTS_BASE_PATH', './documents')

    def create_document(self, task_data: Dict[str, str]) -> str:
        """Crea un nuevo documento con el formato especificado"""
        document = Document()
        self._add_header(document, task_data)
        self._add_content_section(document)
        self._add_references_section(document)

        file_path = self._generate_file_path(task_data)
        document.save(file_path)

        self._open_document(file_path)
        return file_path

    def _add_header(self, document: Document, task_data: Dict[str, str]) -> None:
        """Agrega la portada al documento"""
        # Configuración de la portada
        document.add_heading(task_data['name'], 0)
        document.add_paragraph(f"Asignatura: {task_data['subject']}")
        document.add_paragraph(f"Responsable(s): {task_data['responsible']}")
        document.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")

    def _add_content_section(self, document: Document) -> None:
        """Agrega la sección de contenido"""
        document.add_page_break()
        document.add_heading('Contenido', level=1)
        document.add_paragraph()

    def _add_references_section(self, document: Document) -> None:
        """Agrega la sección de referencias"""
        document.add_page_break()
        document.add_heading('Referencias', level=1)
        document.add_paragraph()

    def _generate_file_path(self, task_data: Dict[str, str]) -> str:
        """Genera la ruta del archivo basada en los datos de la tarea"""
        subject_path = Path(self.base_path) / task_data['subject']
        subject_path.mkdir(parents=True, exist_ok=True)

        filename = f"{task_data['name']}_{task_data['subject']}_{
            task_data['responsible']}.{task_data['file_type']}"
        return str(subject_path / filename)

    def _open_document(self, file_path: str) -> None:
        """Abre el documento con la aplicación predeterminada"""
        os.startfile(file_path) if os.name == 'nt' else os.system(
            f'open "{file_path}"')


# core/document_manager/pdf_converter.py


class PdfAlchemist:
    def transmute_to_pdf(self, source_path: str) -> str:
        """Convierte el documento a PDF"""
        source_path = Path(source_path)
        pdf_path = source_path.with_suffix('.pdf')

        if os.name == 'nt':
            self._convert_windows(source_path, pdf_path)
        else:
            self._convert_unix(source_path, pdf_path)

        return str(pdf_path)

    def _convert_windows(self, source_path: Path, pdf_path: Path) -> None:
        """Convierte a PDF en Windows usando Word COM"""
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(str(source_path))
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()

    def _convert_unix(self, source_path: Path, pdf_path: Path) -> None:
        """Convierte a PDF en sistemas Unix usando LibreOffice"""
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf',
                       str(source_path), '--outdir', str(source_path.parent)])


class AlchemyError(Exception):
    """Excepción personalizada para errores en el proceso de generación/conversión"""
    pass
