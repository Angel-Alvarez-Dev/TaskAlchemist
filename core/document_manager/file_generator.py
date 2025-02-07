import os
from datetime import datetime
from docx import Document

class DocumentAlchemist:
    """
    Clase encargada de generar documentos basados en tareas académicas extraídas de Notion.
    """

    def __init__(self, base_path: str = "documents"):
        """
        Inicializa el generador de documentos.

        Args:
            base_path (str): Carpeta base donde se almacenarán los documentos generados.
        """
        self.base_path = base_path
        # Se asegura de que la carpeta base exista; si no, la crea.
        os.makedirs(self.base_path, exist_ok=True)

    def create_document(self, task_data: dict) -> str:
        """
        Crea un documento a partir de la información de la tarea.

        Args:
            task_data (dict): Diccionario con los datos de la tarea.

        Returns:
            str: Ruta completa del archivo generado.
        """
        # Extrae los datos básicos de la tarea o utiliza valores por defecto.
        name = task_data.get("name", "Untitled")
        subject = task_data.get("subject", "General")
        responsible = task_data.get("responsible", "Unknown")
        file_type = task_data.get("file_type", "docx").lower()

        # Formatea el nombre del archivo combinando el nombre, asignatura y responsable,
        # y reemplaza espacios por guiones bajos para evitar problemas en la ruta.
        file_name = f"{name}_{subject}_{responsible}.{file_type}".replace(" ", "_")
        
        # Define la ruta específica para la asignatura y se asegura de que exista.
        subject_path = os.path.join(self.base_path, subject)
        os.makedirs(subject_path, exist_ok=True)
        
        # Construye la ruta completa del archivo
        file_path = os.path.join(subject_path, file_name)

        # Verifica el tipo de archivo y llama a la función correspondiente para crearlo.
        if file_type == "docx":
            self._create_docx(file_path, task_data)
        else:
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")

        return file_path

    def _create_docx(self, file_path: str, task_data: dict):
        """
        Crea un documento Word (.docx) con una estructura básica que incluye:
        - Un título (nombre de la tarea)
        - Metadatos como asignatura, responsable y fecha de creación
        - Un separador visual y un párrafo inicial para el contenido

        Args:
            file_path (str): Ruta donde se guardará el documento.
            task_data (dict): Datos de la tarea para rellenar el documento.
        """
        # Crea un objeto Document de python-docx
        doc = Document()
        
        # Agrega un encabezado con el nombre de la tarea o un valor por defecto.
        doc.add_heading(task_data.get("name", "Documento"), level=1)

        # Agrega párrafos con metadatos básicos de la tarea.
        doc.add_paragraph(f"Asignatura: {task_data.get('subject', 'General')}")
        doc.add_paragraph(f"Responsable: {task_data.get('responsible', 'Unknown')}")
        doc.add_paragraph(f"Fecha de creación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Agrega un separador visual y un párrafo inicial para el contenido.
        doc.add_paragraph("\n---\n")
        doc.add_paragraph("Aquí comienza el contenido del documento...")

        # Guarda el documento en la ruta especificada.
        doc.save(file_path)
        print(f"📄 Documento creado: {file_path}")
